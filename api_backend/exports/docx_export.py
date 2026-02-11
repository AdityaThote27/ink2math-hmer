from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt

def generate_docx(filename, input_expr, solution, steps):
    doc = Document()

    # ---------- TITLE ----------
    title = doc.add_heading("INK2MATH: SMART EQUATION SOLVER", level=1)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph()

    # ---------- QUESTION ----------
    doc.add_heading("Question:", level=2)
    doc.add_paragraph(input_expr)

    doc.add_paragraph()

    # ---------- SOLUTION ----------
    doc.add_heading("Solution:", level=2)
    doc.add_paragraph(f"x = {solution}")

    doc.add_paragraph()

    # ---------- STEPS ----------
    doc.add_heading("Steps to Solve:", level=2)

    for step in steps:
        doc.add_paragraph(
            f"{step['step']}. {step['description']}"
        )

    doc.save(filename)
